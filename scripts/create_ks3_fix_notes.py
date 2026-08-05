from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path("/Users/stevenchristian/Documents/claverio/P2P/koperasi-ks3/landing/Issue & Fix Notes - Website KS3.docx")


def set_font(run, size=11, color="1F2937", bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_widths(table, widths):
    table.autofit = False
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_text(doc, text, size=11, color="1F2937", bold=False, after=6, before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    set_font(p.add_run(text), size, color, bold)
    return p


def add_section_title(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(15)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.keep_with_next = True
    set_font(p.add_run(title), 14, "1F4D78", True)
    return p


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)
    sec.header_distance = Inches(0.3)
    sec.footer_distance = Inches(0.3)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run("KOPERASI KS3  |  ISSUE & FIX NOTES"), 8.5, "6B7280", True)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("Dokumen kerja — lengkapi bukti screenshot dan jawaban sebelum dikirim."), 8.5, "6B7280")

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    set_font(title.add_run("Issue & Fix Notes"), 23, "0B2545", True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    set_font(subtitle.add_run("Website Koperasi KS3"), 13, "2E74B5", True)

    meta = doc.add_table(rows=3, cols=2)
    meta.style = "Table Grid"
    set_table_widths(meta, [1800, 7560])
    metadata = [("Tanggal", "1 Agustus 2026"), ("Tujuan", "Rekap perbaikan, konfirmasi, dan bukti implementasi"), ("Status dokumen", "Berjalan — siap dilengkapi screenshot dan jawaban")]
    for row, (label, value) in zip(meta.rows, metadata):
        shade(row.cells[0], "E8EEF5")
        for i, value_text in enumerate((label, value)):
            p = row.cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_font(p.add_run(value_text), 9.5, "1F2937", i == 0)

    add_text(doc, "Cara pakai: setiap item dapat diberi status, dilampiri screenshot sebelum/sesudah, lalu pertanyaan yang masih terbuka diisi pada kolom catatan/jawaban.", 9.5, "4B5563", after=8, before=10)

    items = [
        ("Website — Header", "Hapus alamat kantor pada bagian kanan atas header.", "Perbaikan", "Selesai", "Elemen alamat sudah dihapus dari header.", "Lampirkan screenshot header terbaru."),
        ("Website — Header", "Hapus tulisan dan indikator hijau “Koperasi KS3” pada sisi kiri atas header.", "Perbaikan", "Selesai", "Indikator dan teks sudah dihapus dari header.", "Lampirkan screenshot header terbaru."),
        ("Website — Kontak", "Logo WhatsApp di bagian pertanyaan: apakah dapat diarahkan ke chat WhatsApp dari nomor yang dicantumkan?", "Pertanyaan", "Butuh jawaban", "Konfirmasi nomor WhatsApp dan format pesan pembuka bila diperlukan.", "Jawaban/keputusan: ______________________________"),
        ("Website — Home", "Cek tautan logo backend dan store pada halaman Home; diduga link Play Store dan Apple App Store tertukar.", "Bug", "Perlu dicek", "Verifikasi tujuan URL masing-masing tombol/logo.", "Bukti: screenshot atau link tujuan setelah diperbaiki."),
        ("Website — Home", "Tutup/sembunyikan logo Apple App Store karena aplikasi belum tersedia di App Store.", "Perbaikan", "Menunggu konfirmasi", "Tentukan apakah disembunyikan penuh atau diberi label “Segera hadir”.", "Jawaban/keputusan: ______________________________"),
        ("Website — Kontak", "Konfirmasi apakah website perlu mendukung beberapa alamat cabang.", "Pertanyaan", "Butuh jawaban", "Jika ya, mohon daftar nama cabang, alamat, kontak, dan prioritas tampilannya.", "Jawaban/keputusan: ______________________________"),
        ("Website — Footer/Sosial", "Perbarui logo Instagram.", "Perbaikan", "Menunggu aset", "Mohon kirim file logo/ikon Instagram yang ingin digunakan atau referensi desainnya.", "Bukti: screenshot setelah logo diperbarui."),
        ("Website — Browser tab", "Judul tab/browser masih menampilkan “Peer to Peer”.", "Bug", "Perlu diperbaiki", "Ganti title/meta title agar sesuai identitas Website Koperasi KS3.", "Bukti: screenshot tab browser setelah diperbarui."),
        ("Website — Footer/Legal", "Kebijakan Privasi dan Syarat & Ketentuan: kontennya akan diperbarui dari mana?", "Pertanyaan", "Butuh jawaban", "Konfirmasi apakah konten dikelola melalui CMS atau perlu disiapkan halaman khusus.", "Jawaban/keputusan: ______________________________"),
        ("Payment Gateway", "Aktifkan hanya metode pembayaran transfer bank.", "Konfigurasi", "Butuh konfirmasi", "Konfirmasi rekening/kanal transfer yang aktif serta metode lain yang harus dinonaktifkan.", "Jawaban/keputusan: ______________________________"),
        ("Payment Gateway", "Pastikan nominal biaya admin yang berlaku.", "Pertanyaan", "Butuh jawaban", "Mohon konfirmasi nominal, pihak pembebanan, dan apakah biaya berbeda per metode.", "Jawaban/keputusan: ______________________________"),
        ("CMS — Halaman", "Konfirmasi lokasi untuk mengubah konten halaman pada backend/CMS.", "Panduan", "Perlu dijelaskan", "Siapkan panduan singkat menu CMS untuk pembaruan konten halaman.", "Catatan/panduan: _________________________________"),
        ("CMS — Homepage Hero", "Field “Hero main image alt*” muncul di bagian mana?", "Pertanyaan", "Perlu dijelaskan", "Jelaskan bahwa field ini adalah alt text untuk aksesibilitas/SEO dan tampil saat gambar tidak dimuat atau dibaca screen reader.", "Jawaban: ________________________________________"),
        ("CMS — Homepage Advantage", "Field “Advantages image alt*” muncul di bagian mana?", "Pertanyaan", "Perlu dijelaskan", "Jelaskan bahwa field ini adalah alt text untuk gambar pada section Advantages, bukan teks dekoratif utama.", "Jawaban: ________________________________________"),
    ]

    add_section_title(doc, "A. Website dan Konten")
    website_items = items[:9]
    gateway_items = items[9:11]
    cms_items = items[11:]

    def add_item(no, item):
        area, request, kind, status, note, evidence = item
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        set_table_widths(table, [1650, 7710])
        shade(table.cell(0, 0), "E8EEF5")
        left = table.cell(0, 0).paragraphs[0]
        left.paragraph_format.space_after = Pt(1)
        set_font(left.add_run(f"{no:02d}"), 18, "1F4D78", True)
        p = table.cell(0, 0).add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(area), 8.5, "4B5563", True)
        right = table.cell(0, 1).paragraphs[0]
        right.paragraph_format.space_after = Pt(2)
        set_font(right.add_run(request), 10.2, "1F2937", True)
        for label, value in (("Tipe", kind), ("Status", status), ("Catatan", note), ("Bukti / jawaban", evidence)):
            p = table.cell(0, 1).add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            set_font(p.add_run(f"{label}: "), 9, "374151", True)
            set_font(p.add_run(value), 9, "374151")
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    for number, item in enumerate(website_items, 1):
        add_item(number, item)
    add_section_title(doc, "B. Payment Gateway")
    for number, item in enumerate(gateway_items, 10):
        add_item(number, item)
    add_section_title(doc, "C. Backend / CMS")
    for number, item in enumerate(cms_items, 12):
        add_item(number, item)

    add_section_title(doc, "Ringkasan untuk Pengiriman")
    add_text(doc, "Status saat ini: 2 perbaikan header telah selesai. Item lain menunggu pengecekan, aset, atau keputusan. Sertakan screenshot pada item yang sudah selesai dan isi jawaban pada item bertanda “Butuh jawaban” agar pembaruan berikutnya dapat langsung dikerjakan.", 10, "374151", after=0)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
